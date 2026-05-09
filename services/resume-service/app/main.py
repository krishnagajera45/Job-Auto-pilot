from __future__ import annotations

import os
import shutil
import subprocess
import uuid
from pathlib import Path
from typing import Dict, List, Literal, Optional

from fastapi import FastAPI, HTTPException
from docx import Document
from pydantic import BaseModel, Field

app = FastAPI(title="Job Autopilot Resume Service", version="0.1.0")

# Prototype-only in-memory storage; replace with Postgres + object storage.
RESUMES: Dict[str, dict] = {}
COVER_LETTERS: Dict[str, dict] = {}

TEMPLATE_DIR = Path(__file__).parent / "templates"
OUTPUT_DIR = Path(os.getenv("DOCUMENT_OUTPUT_DIR", "/tmp/job-autopilot/documents"))


class ResumeCreateRequest(BaseModel):
    user_id: str
    title: str
    content: str
    source: Literal["upload", "linkedin", "paste"] = Field(
        default="upload",
        description="upload | linkedin | paste",
    )


class ResumeVersionRequest(BaseModel):
    content: str
    summary: Optional[str] = None


class CoverLetterRequest(BaseModel):
    user_id: str
    job_id: str
    content: str


class RenderDocumentsRequest(BaseModel):
    user_id: str
    job_id: str
    resume_content: str
    cover_letter_content: str
    name: str = "Candidate"
    job_title: str = "Role"
    company: str = "Company"
    template_name: Literal["default"] = "default"


@app.get("/health")
async def health_check() -> dict:
    return {"status": "ok", "service": "resume-service"}


@app.post("/v1/resumes")
async def create_resume(request: ResumeCreateRequest) -> dict:
    resume_id = str(uuid.uuid4())
    RESUMES[resume_id] = {
        "id": resume_id,
        "user_id": request.user_id,
        "title": request.title,
        "versions": [
            {
                "version_id": str(uuid.uuid4()),
                "content": request.content,
                "summary": "Initial version",
            }
        ],
    }
    return RESUMES[resume_id]


@app.get("/v1/resumes")
async def list_resumes(user_id: str) -> List[dict]:
    return [resume for resume in RESUMES.values() if resume["user_id"] == user_id]


@app.post("/v1/resumes/{resume_id}/versions")
async def add_version(resume_id: str, request: ResumeVersionRequest) -> dict:
    resume = RESUMES.get(resume_id)
    if not resume:
        raise HTTPException(status_code=404, detail="Resume not found")
    version = {
        "version_id": str(uuid.uuid4()),
        "content": request.content,
        "summary": request.summary or "Tailored version",
    }
    resume["versions"].append(version)
    return {"resume_id": resume_id, "version": version}


@app.post("/v1/cover-letters")
async def create_cover_letter(request: CoverLetterRequest) -> dict:
    cover_id = str(uuid.uuid4())
    COVER_LETTERS[cover_id] = {
        "id": cover_id,
        "user_id": request.user_id,
        "job_id": request.job_id,
        "content": request.content,
    }
    return COVER_LETTERS[cover_id]


def render_latex(template_path: Path, values: dict) -> str:
    if not template_path.exists():
        raise HTTPException(status_code=404, detail=f"Template not found: {template_path.name}")
    template_text = template_path.read_text()
    for key, value in values.items():
        template_text = template_text.replace(f"{{{{{key}}}}}", str(value))
    return template_text


def compile_pdf(tex_path: Path) -> tuple[Optional[Path], Optional[Path]]:
    pdflatex = shutil.which("pdflatex")
    if not pdflatex:
        return None, None
    result = subprocess.run(
        [pdflatex, "-interaction=nonstopmode", tex_path.name],
        cwd=tex_path.parent,
        check=False,
        stdout=subprocess.PIPE,
        stderr=subprocess.STDOUT,
        text=True,
    )
    log_path = tex_path.with_suffix(".log")
    log_path.write_text(result.stdout or "")
    pdf_path = tex_path.with_suffix(".pdf")
    return (pdf_path if pdf_path.exists() else None, log_path)


def write_docx(docx_path: Path, heading: str, content: str) -> None:
    document = Document()
    document.add_heading(heading, level=1)
    for line in content.splitlines():
        if line.strip():
            document.add_paragraph(line)
    document.save(docx_path)


@app.post("/v1/documents/render")
async def render_documents(request: RenderDocumentsRequest) -> dict:
    bundle_id = str(uuid.uuid4())
    bundle_dir = OUTPUT_DIR / bundle_id
    bundle_dir.mkdir(parents=True, exist_ok=True)

    values = {
        "name": request.name,
        "job_title": request.job_title,
        "company": request.company,
        "content": request.resume_content,
    }
    resume_template = TEMPLATE_DIR / "resume.tex"
    cover_template = TEMPLATE_DIR / "cover_letter.tex"

    resume_tex_path = bundle_dir / "resume.tex"
    cover_tex_path = bundle_dir / "cover_letter.tex"

    resume_tex_path.write_text(render_latex(resume_template, values))
    cover_tex_path.write_text(
        render_latex(
            cover_template,
            {
                "name": request.name,
                "job_title": request.job_title,
                "company": request.company,
                "content": request.cover_letter_content,
            },
        )
    )

    resume_pdf, resume_log = compile_pdf(resume_tex_path)
    cover_pdf, cover_log = compile_pdf(cover_tex_path)

    resume_docx = bundle_dir / "resume.docx"
    cover_docx = bundle_dir / "cover_letter.docx"
    write_docx(resume_docx, "Tailored Resume", request.resume_content)
    write_docx(cover_docx, "Cover Letter", request.cover_letter_content)

    return {
        "bundle_id": bundle_id,
        "output_dir": str(bundle_dir),
        "resume": {
            "tex_path": str(resume_tex_path),
            "pdf_path": str(resume_pdf) if resume_pdf else None,
            "docx_path": str(resume_docx),
            "latex_log": str(resume_log) if resume_log else None,
        },
        "cover_letter": {
            "tex_path": str(cover_tex_path),
            "pdf_path": str(cover_pdf) if cover_pdf else None,
            "docx_path": str(cover_docx),
            "latex_log": str(cover_log) if cover_log else None,
        },
        "status": "rendered" if resume_pdf and cover_pdf else "tex-only",
    }
